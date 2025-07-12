import os
import asyncio
from typing import Optional, Dict, Any, List
from pypdf import PdfReader
from docx import Document
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
import csv
import json

# Увеличил максимальный размер, так как современные модели могут обрабатывать больше текста
MAX_TEXT_SIZE = 150_000  # Maximum number of characters for a single file

# Словарь с поддерживаемыми форматами для быстрой проверки
SUPPORTED_FORMATS = {
    ".pdf", ".docx", ".doc", ".txt", ".md", ".rtf", ".odt", 
    ".html", ".htm", ".csv", ".json", ".py", ".js", ".css"
}

def extract_text_from_pdf(path: str) -> str:
    """Извлекает текст из PDF файла."""
    try:
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        text = text.strip()
        if text:
            return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
        return "[PDF is empty or unreadable.]"
    except Exception as e:
        print(f"Error reading PDF ({os.path.basename(path)}): {e}")
        return f"[Error reading PDF ({os.path.basename(path)}): {e}]"

def extract_text_from_docx(path: str) -> str:
    """Извлекает текст из DOCX файла."""
    try:
        doc = Document(path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Обрабатываем таблицы, если они есть
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        text = "\n".join(paragraphs).strip()
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except Exception as e:
        print(f"Error reading DOCX ({os.path.basename(path)}): {e}")
        return f"[Error reading DOCX ({os.path.basename(path)}): {e}]"

def extract_text_from_txt(path: str) -> str:
    """Извлекает текст из TXT файла."""
    encodings = ["utf-8", "latin-1", "cp1251", "ascii"]  # Пробуем разные кодировки
    
    for encoding in encodings:
        try:
            with open(path, encoding=encoding) as f:
                text = f.read()
            return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"Error reading TXT ({os.path.basename(path)}): {e}")
            return f"[Error reading TXT ({os.path.basename(path)}): {e}]"
    
    return f"[Error: Unable to decode file ({os.path.basename(path)}) with available encodings]"

def extract_text_from_rtf(path: str) -> str:
    """Извлекает текст из RTF файла."""
    try:
        with open(path, 'rb') as f:
            content = f.read().decode('latin-1')
            text = rtf_to_text(content)
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except Exception as e:
        print(f"Error reading RTF ({os.path.basename(path)}): {e}")
        return f"[Error reading RTF ({os.path.basename(path)}): {e}]"

def extract_text_from_odt(path: str) -> str:
    """Извлекает текст из ODT файла."""
    try:
        # Try odfpy first; fallback to textract if not available
        try:
            from odf.opendocument import load
            from odf.text import P
            doc = load(path)
            text = "\n".join([str(t) for t in doc.getElementsByType(P)])
            text = text.strip()
        except ImportError:
            try:
                import textract
                text = textract.process(path).decode("utf-8")
            except Exception as e:
                print(f"Error reading ODT ({os.path.basename(path)}): {e}")
                return f"[Error reading ODT ({os.path.basename(path)}): {e}]"
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except Exception as e:
        print(f"Error reading ODT ({os.path.basename(path)}): {e}")
        return f"[Error reading ODT ({os.path.basename(path)}): {e}]"

def extract_text_from_html(path: str) -> str:
    """Извлекает текст из HTML файла."""
    try:
        with open(path, encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            # Удаляем скрипты, стили и другой нетекстовый контент
            for tag in soup(['script', 'style', 'meta', 'link']):
                tag.decompose()
            text = soup.get_text(separator="\n")
            # Удаляем лишние пустые строки
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            text = "\n".join(lines)
        text = text.strip()
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except Exception as e:
        print(f"Error reading HTML ({os.path.basename(path)}): {e}")
        return f"[Error reading HTML ({os.path.basename(path)}): {e}]"

def extract_text_from_csv(path: str) -> str:
    """Извлекает текст из CSV файла."""
    try:
        result = []
        # Пробуем разные разделители
        for delimiter in [',', ';', '\t']:
            try:
                with open(path, encoding="utf-8", newline='') as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    rows = list(reader)
                    if len(rows) > 0 and len(rows[0]) > 1:  # Если нашли подходящий разделитель
                        for row in rows:
                            result.append(" | ".join(cell for cell in row if cell))
                        break
            except Exception:
                continue
        
        if not result:  # Если ни один разделитель не подошел
            with open(path, encoding="utf-8") as f:
                result = f.readlines()
        
        text = "\n".join(result).strip()
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except Exception as e:
        print(f"Error reading CSV ({os.path.basename(path)}): {e}")
        return f"[Error reading CSV ({os.path.basename(path)}): {e}]"

def extract_text_from_json(path: str) -> str:
    """Извлекает текст из JSON файла и форматирует его."""
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        text = json.dumps(data, indent=2, ensure_ascii=False)
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except Exception as e:
        print(f"Error reading JSON ({os.path.basename(path)}): {e}")
        return f"[Error reading JSON ({os.path.basename(path)}): {e}]"

def extract_text_from_code(path: str) -> str:
    """Извлекает текст из файла с кодом (Python, JS, и т.д.)."""
    return extract_text_from_txt(path)

def extract_text_from_md(path: str) -> str:
    """Извлекает текст из Markdown файла."""
    # Treat markdown as plain text
    return extract_text_from_txt(path)

def extract_text_from_doc(path: str) -> str:
    """Извлекает текст из DOC файла."""
    # Try using textract for legacy .doc files
    try:
        import textract
        text = textract.process(path).decode("utf-8")
        text = text.strip()
        return text[:MAX_TEXT_SIZE] + ('\n[Truncated]' if len(text) > MAX_TEXT_SIZE else '')
    except ImportError:
        return f"[Error: textract module not available for reading DOC files]"
    except Exception as e:
        print(f"Error reading DOC ({os.path.basename(path)}): {e}")
        return f"[Error reading DOC ({os.path.basename(path)}): {e}]"

def extract_text_from_file(path: str) -> str:
    """Извлекает текст из файла в зависимости от его типа."""
    # Проверка существования файла
    if not os.path.exists(path):
        return f"[Error: File not found: {path}]"
    
    ext = os.path.splitext(path)[-1].lower()
    
    # Проверка поддерживаемого формата
    if ext not in SUPPORTED_FORMATS:
        return f"[Unsupported file type: {os.path.basename(path)}]"
    
    # Вызов соответствующей функции в зависимости от типа файла
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".doc":
        return extract_text_from_doc(path)
    elif ext == ".txt":
        return extract_text_from_txt(path)
    elif ext == ".md":
        return extract_text_from_md(path)
    elif ext == ".rtf":
        return extract_text_from_rtf(path)
    elif ext == ".odt":
        return extract_text_from_odt(path)
    elif ext in [".html", ".htm"]:
        return extract_text_from_html(path)
    elif ext == ".csv":
        return extract_text_from_csv(path)
    elif ext == ".json":
        return extract_text_from_json(path)
    elif ext in [".py", ".js", ".css"]:
        return extract_text_from_code(path)
    else:
        return f"[Unsupported file type: {os.path.basename(path)}]"

async def extract_text_from_file_async(path: str) -> str:
    """Асинхронно извлекает текст из файла."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, extract_text_from_file, path)

async def save_file_async(file_path: str, content: str) -> bool:
    """
    Асинхронно сохраняет содержимое в файл.
    Создает директории, если их нет.
    """
    try:
        # Создаем директории, если их нет
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # Асинхронно сохраняем содержимое
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, lambda: _save_file(file_path, content))
        
        return True
    except Exception as e:
        print(f"Error saving file: {e}")
        return False

def _save_file(file_path: str, content: str) -> None:
    """Вспомогательная функция для сохранения файла."""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

async def list_files_async(directory: str, pattern: Optional[str] = None) -> List[str]:
    """
    Асинхронно получает список файлов в директории, 
    опционально фильтруя по паттерну.
    """
    try:
        loop = asyncio.get_event_loop()
        files = await loop.run_in_executor(None, os.listdir, directory)
        
        if pattern:
            import fnmatch
            files = [f for f in files if fnmatch.fnmatch(f, pattern)]
        
        # Возвращаем полные пути к файлам
        return [os.path.join(directory, f) for f in files]
    except Exception as e:
        print(f"Error listing files: {e}")
        return []
